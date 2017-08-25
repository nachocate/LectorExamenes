from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import csv
import sys
import unicodedata
import codecs

reload(sys)  # Reload does the trick!
sys.setdefaultencoding('UTF8')

def leerPreguntas(nombre):
	pregunta=[]
	rta1=[]
	rta2=[]
	rta3=[]
	rta4=[]
	rta5=[]
	with open(nombre, 'rb') as csvfile:
		vector = csv.reader(csvfile, delimiter=';', quotechar='|')
		for i in vector:
			pregunta.append(str(i[0]))
			rta1.append(str(i[1]))
			rta2.append(str(i[2]))
			rta3.append(str(i[3]))
			rta4.append(str(i[4]))
			rta5.append(str(i[5]))
		csvfile.close()
	return pregunta,rta1,rta2,rta3,rta4,rta5	

def leerPreguntasFijas(nombre):
	Mat=[];
	#file = open(nombre,'r')
	file = codecs.open(nombre, encoding='utf-8')
    #print repr(line)
	cont=0
	vec=[]
	for line in file:
		Mat.append(line)
	file.close()
	return Mat



def generarExamen(nombre,M):
	sec=["a)","b)","c)","d)","e)"]
	document = Document()
	p = document.add_paragraph('Examen Final ')
	for i in range(0,len(M)):
		preg=M[i]
		print i,preg,len(M),M[1]
		p=document.add_paragraph()
		p=document.add_paragraph()
		linea=str(i+1)+") "+preg[0]
		p.add_run(linea).font.size=Pt(10)
		for j in range(1,len(preg)):
			p=document.add_paragraph()
			ff=p.add_run(sec[j-1]+" "+preg[j])
			ff.font.size=Pt(10)
			ff.font.name="Arial"
	document.save(nombre+".docx")
	

def generarExamenPlano(nombre,M):
	sec=["a)","b)","c)","d)","e)"]
	document = Document()
	p = document.add_paragraph('Examen Final ')
	cont_preg=1;
	cont=0
	for i in M:
		if cont==0:
			p=document.add_paragraph()
			p=document.add_paragraph()
			linea=str(cont_preg)+") "+i
			cont_preg=cont_preg+1
			p.add_run(linea).font.size=Pt(10)
			cont=cont+1
		else:
			
			p=document.add_paragraph()
			ff=p.add_run(sec[cont-1]+" "+i)
			ff.font.size=Pt(10)
			ff.font.name="Arial"
			cont=cont+1
			if cont==6:
				cont=0
			
	document.save(nombre+".docx")
	










Preg=["Cual es la cosas de la pregunta 1 que solo tiene cosas de la cosas de nada? solo contemple el echo de que coso y ademas que coso y las siguientes preguntas","opcion 1 con veneno","solo comer papas fritas","tiroides","anemia","falso"]

Preg2=["Cuanto es el resultado de hacer 2 +2?","4","5","16","8","32"]
Preg1=["Cuanto es 3x4","1","2"]

M=[]
M.append(Preg)
M.append(Preg1)
M.append(Preg2)
M.append(Preg)
M.append(Preg1)
M.append(Preg2)
M.append(Preg)
M.append(Preg1)
M.append(Preg2)
M.append(Preg)
M.append(Preg1)
M.append(Preg2)

#generarExamen("examen",M)


#print leerPreguntas("preguntas.csv")
Mat=leerPreguntasFijas("preguntas1.txt")
for i in Mat:
	print i

generarExamenPlano("examenEcho",Mat)
#generarExamen("examenEcho",Mat)
