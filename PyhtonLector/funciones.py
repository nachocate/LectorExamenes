# import the necessary packages
import numpy as np
import cv2
import math as m
import csv
import Alumno as alu
import sys
import unicodedata
import codecs
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import os

reload(sys)  # Reload does the trick!
sys.setdefaultencoding('UTF8')

def leerPreguntas(nombre):
	pregunta=[]
	rta1=[]
	rta2=[]
	rta3=[]
	rta4=[]
	rta5=[]
	with open(nombre, 'rb') as csvfile:
		vector = csv.reader(csvfile, delimiter=';', quotechar='|')
		for i in vector:
			pregunta.append(str(i[0]))
			rta1.append(str(i[1]))
			rta2.append(str(i[2]))
			rta3.append(str(i[3]))
			rta4.append(str(i[4]))
			rta5.append(str(i[5]))
		csvfile.close()
	return pregunta,rta1,rta2,rta3,rta4,rta5	

def leerPreguntasFijas(nombre):
	Mat=[];
	#file = open(nombre,'r')
	file = codecs.open(nombre, encoding='utf-8')
    #print repr(line)
	cont=0
	vec=[]
	for line in file:
		Mat.append(line)
	file.close()
	return Mat

def generarExamenPlano(nombre,M):
	sec=["a)","b)","c)","d)","e)"]
	document = Document()
	p = document.add_paragraph('Examen Final ')
	cont_preg=1;
	cont_rta=0
	print  "Info de la M"
	print "Tamanio de M",len(M)
	for i in M:
		print "Por cada i de M", len(i)
		flag=True
		cont_rta=0
		for j in i:
			if flag:
				p=document.add_paragraph()
				p=document.add_paragraph()
				linea=str(cont_preg)+") "+j
				p.add_run(linea).font.size=Pt(10)
				flag=False
			else:
				p=document.add_paragraph()
				ff=p.add_run(sec[cont_rta]+" "+j)
				ff.font.size=Pt(10)
				ff.font.name="Arial"
				cont_rta=cont_rta+1
		cont_preg=cont_preg+1
			
	document.save(nombre+".docx")
	






def exist(ids,vec):
	for i in vec:
		if i==ids:
			return True
	return False

def rest(a,b):
	return (a[0]-b[0],a[1]-b[1])

def suma(a,b):
	return (a[0]+b[0],a[1]+b[1])


def getPoint(a,b,c):
	return (a[0] +c*b[0],a[1] +c*b[1])

def getUni(a):
	norma=(1/cv2.norm(a))
	return(a[0]*norma,a[1]*norma)

def pp(a,b):
	return a[0]*b[0]+a[1]*b[1]

def getVec(p1, p2):
	return rest(p2,p1)

def getVecUni(p1,p2):
	return getUni(getVec(p1,p2))

def CorreccionAngulos(Ref1,Ref2):
	vNormal=(0,1)
	return m.acos(pp(getUni(rest(Ref2,Ref1)),vNormal))

def getAng(a,b):
	return m.acos(pp(getUni(a),getUni(b)))

def CorregirExamen(path):
	img=cv2.imread(path)
	alto, ancho, canales = img.shape;
	img=cv2.resize(img,(500, 800), interpolation = cv2.INTER_CUBIC)
	
	alto, ancho, canales = img.shape;
	img_g = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
	ret,img_g = cv2.threshold(img_g,200,255,cv2.THRESH_BINARY_INV)

	cv2.imwrite('/home/nacho/Escritorio/PyhtonLector/binaria.png',img_g)

	#kernel = np.ones((3,3))
	#img_g = cv2.erode(img_g,kernel,3)
	#ero = cv2.erode(img_g,kernel,20)
	#img_g = cv2.dilate(img_g,kernel,1)

	kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (3,3))
	#im = np.zeros((100,100), dtype=np.uint8)
	#im[50:,50:] = 255
	ero = cv2.erode(img_g, kernel, iterations = 3)
	dil = cv2.dilate(ero, kernel, iterations = 2)
	
	cv2.imwrite('/home/nacho/Escritorio/PyhtonLector/morpo.png',img_g)
	
	cv2.imwrite('/home/nacho/Escritorio/PyhtonLector/erode.png',ero)
	cv2.imwrite('/home/nacho/Escritorio/PyhtonLector/dilate.png',dil)

	img_g=dil
#--------------------------------------------------------------------------
	
	
	cant_v,labels,stats,cen = cv2.connectedComponentsWithStats(img_g, 8, cv2.CV_32S)
	#print cant_v
	
	#print cen[2,]
	Ref=(cen[2,0]-cen[1,0],cen[2,1]-cen[1,1])
	Ref3=(cen[3,0]-cen[1,0],cen[3,1]-cen[1,1])
	Ref4=(cen[4,0]-cen[1,0],cen[4,1]-cen[1,1])
	Ref8=(cen[8,0]-cen[1,0],cen[8,1]-cen[1,1])
	
	RR3= (getAng(Ref,Ref3))*180/m.pi
	RR4= (getAng(Ref,Ref4))*180/m.pi
	vHorizontal=(1.0,1.0)
	vVertical=getUni(Ref8)
	if(m.fabs(RR3-RR4)<2):
		vHorizontal=getUni(Ref)
	else:
		vHorizontal=getUni(suma(Ref3,Ref4))
	#	/**
	#	Aqui lo que recorre de los nodos, son los 2 a 7, osea 6 
	#	primeros cuadraditos negros
	#	
	#	**/
	dist=[]
	resultado=[]
	
	for i in range(2,8):
		vaux=(cen[i, 0]-cen[1,0], cen[i, 1]-cen[1,1])
		dist.append(cv2.norm(vaux))
		img=cv2.circle(img, (int(cen[i,0]),int(cen[i,1])), 3, (0,0,255))
		
	
	#/** Funcion que calcula angulos**/
	ang=CorreccionAngulos((cen[1, 0],cen[1, 1]),(cen[8, 0],cen[8, 1]))
	#	cout<<"------   ANGULO IMAGEN: "<<ang*180/CV_PI<<endl
	
	#/** PRIMER CUADRADITO IZQUIERDO INFERIOR **/
	

	#/** Aqui detecto el examen **/
	P8=(cen[8, 0], cen[8, 1])
	contador=0
	idex=0
	for i in range(0,12):
		pExam=getPoint(P8,vHorizontal,dist[0]*(1+i*0.0666))
		img=cv2.circle(img, (int(pExam[0]),int(pExam[1])), 3, (255,0,255))
		if(img_g.item(int(pExam[1]),int(pExam[0]))>200):
			contador=contador+1
			idex=idex+m.pow(2,12-i-1)
	
	#cout<<"--------Num EXAMEN---------------------"<<endl
	#print "contador para examen: "+ str(contador)
	#print "idex: "+ str(idex)
	#print img_g.item(339,159),img_g[339,159]
	#print img_g[159,341],img_g.item(159,341)
	
	
		
	#/** Aqui detecto el id del alumno **/
	indi=9+contador
	PINDI=(cen[indi, 0], cen[indi, 1])
	contador=0
	idal=0
	for i in range(0,12):
		pAlu=getPoint(PINDI,vHorizontal,dist[0]*(1+i*0.0666))
		img=cv2.circle(img, (int(pAlu[0]),int(pAlu[1])), 3, (255,0,255))
		if(img_g.item(int(pAlu[1]),int(pAlu[0]))>200):
			contador=contador+1
			idal=idal+m.pow(2,12-i-1)
	#cout<<"--------Num ALUMNO---------------------"<<endl
	#print "contador para alumno: "+ str(contador)
	#print "idAlu: "+ str(idal)
	
	
	
	#/** CORRECCION DEL EXAMEN  **/
	indi1=indi+contador+1
	tam=len(cen)
	#cout<<"Centroides: "<<centroids.rows<<" indi1:"<<indi1<<endl
	for i in range(indi1, tam, 2):
	#for i in range(indi1i:2:tam):
		Pi=(cen[i,0],cen[i,1])
		Pf=(cen[i+1,0],cen[i+1,1])
		img=cv2.circle(img, (int(Pi[0]),int(Pi[1])), 3, (255,255,0))
		if(cv2.norm(getVec(Pi,Pf))<cv2.norm(Ref8)+10):
			i=i-1
			resultado.append(0)
			continue
		img=cv2.circle(img, (int(Pf[0]),int(Pf[1])), 3, (255,255,0))
		nor=cv2.norm(getVec(Pi,Pf))
		mini=6
		dif=10000
		for j in range(0,len(dist)):
			if(m.fabs(nor-dist[j])<dif):
				dif=m.fabs(nor-dist[j])
				mini=j+1
		resultado.append(mini)

	#for i in range(0,len(resultado)):
		#print(resultado[i])

	cv2.imwrite('/home/nacho/Escritorio/PyhtonLector/Circulos.png',img)
	return [int(idal),int(idex),resultado]


def loadCSV(nombre):
	nombres=[]
	apellidos=[]
	dnis=[]
	with open(nombre, 'rb') as csvfile:
		vector = csv.reader(csvfile, delimiter=',', quotechar='|')
		for i in vector:
			nombres.append(str(i[0]))
			apellidos.append(str(i[1]))
			dnis.append(str(i[2]))
		csvfile.close()
	return nombres,apellidos,dnis	
def loadCSVCategoria(nombre):
	nombres=[]
	with open(nombre, 'rb') as csvfile:
		vector = csv.reader(csvfile, delimiter=',', quotechar='|')
		for i in vector:
			nombres.append(str(i[0]))
		csvfile.close()
	return nombres	


#def saveCSV(nombre,Alum):
#	ofile  = open(nombre, "wb")
#	writer = csv.writer(ofile, delimiter=',', quotechar='"', quoting=csv.QUOTE_ALL)
#	for i in Alum:
#		writer.writerow([i.getId(),i.getNombre(),i.getApellido(),i.getDni(),i.getNota()])
#	ofile.close()
	
def saveCSV(nombre,a1,a2,a3,a4,a5):
	ofile  = open(nombre, "wb")
	writer = csv.writer(ofile, delimiter=',', quotechar='"', quoting=csv.QUOTE_ALL)
	for i in range(0,len(a1)):
		writer.writerow([a1[i],a2[i],a3[i],a4[i],a5[i]])
	ofile.close()
def genExamen(path,nombre, apellido,dni,idAlumno,idExamen,pathAux):
	img=cv2.imread(path)
	alto, ancho, canales = img.shape;
	img=cv2.resize(img,(500, 800), interpolation = cv2.INTER_CUBIC)
	alto, ancho, canales = img.shape;
	img_g = cv2.cvtColor(img,cv2.COLOR_RGB2GRAY)
	ret,img_g = cv2.threshold(img_g,200,255,cv2.THRESH_BINARY_INV)
	font = cv2.FONT_HERSHEY_SIMPLEX
	cv2.putText(img,"Nombre y Apellido: "+nombre+" "+apellido,(30,20),font,0.3,(0,0,0),1,1)
	cv2.putText(img,"Dni: "+dni,(30,35), font,0.3,(0,0,0),1,1)
	cv2.putText(img,"Hoja: "+str(idExamen),(30,50), font,0.3,(0,0,0),1,1)
	alt=109
	despIni=270
	inc=15.8
	l=[]
	
	idAlumnoBinario=bin(idAlumno)[2:].zfill(12)
	idExamenBinario=bin(idExamen)[2:].zfill(12)

	
	
	for i in range(0,12):
		if idExamenBinario[i]=='1':
			l.append([despIni+inc*i,alt])
	for k in l:
		x=k[1]
		y=k[0]
		for i in range(0,12):
			for j in range(0,11):
				img[int(x+i),int(y+j)]=[0,0,0]
	
	
	alt=151
	l=[]
	for i in range(0,12):
		if idAlumnoBinario[i]=='1':
			l.append([despIni+inc*i,alt])
	for k in l:
		x=k[1]
		y=k[0]
		for i in range(0,12):
			for j in range(0,11):
				img[int(x+i),int(y+j)]=[0,0,0]
	
	inc=15
	l=[]
	l=range(20*(idExamen-1)+1,(20*(idExamen-1)+21))
	for i in range(0,20):
		cv2.putText(img,str(l[i])+")",(15,int(203+ 28.4*i)), font,0.3,(0,0,0),1,1)
	cv2.imwrite(pathAux+"/"+nombre+apellido+str(idExamen)+'.png',img)

